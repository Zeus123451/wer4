# # import docx

# # mydoc = docx.Document('')
# # mydoc.add_heading("This is level 1 heading", 0)
# # mydoc.add_heading("This is level 2 heading", 1)
# # mydoc.add_heading("This is level 3 heading", 2)
# # mydoc.save("E:/my_written_file.docx")

# import os

# # # Retrieve current working directory (`cwd`)
# # cwd = os.getcwd()
# # cwd

# # # Change directory 
# # os.chdir("C:/Users/Daler Hamidov/Desktop/tesstauto")

# # # List all files and directories in current directory
# # os.listdir('.')
import xlrd
from docx import Document
from docx.shared import Inches

# # Open the Workbook
workbook = xlrd.open_workbook("myex1.xls")
# # Open the 1stworksheet
worksheet = workbook.sheet_by_index(0)
# # Iterate the rows and columns
# # for i in range(0, 5):
# #     for j in range(0, 3):
# # #         # Print the cell values with tab space
# # #         print(worksheet.cell_value(1, 2), end='\n')
# # #     print('\n')
# # print(worksheet.cell_value(0, 1), end='\n')
# document = Document()

# # document.add_heading('Document Title', 0)

# # p = document.add_paragraph(worksheet.cell_value(1, 2))
# # p.add_run('bold').bold = True
# # p.add_run(' and some ')
# # p.add_run('italic.').italic = True
# # for quan_of_rows in range(0,200):
# # 	for column in range(0,5):
# # 		# print(quan_of_rows,column)
# # 		p = document.add_paragraph(worksheet.cell_value(quan_of_rows, column))
# # print(p.add_run(worksheet.cell_value(1, 1)).font.highlight_color)
# from docx.enum.text import WD_COLOR_INDEX
# number_q = 1
# for quan_of_rows in range(0,500):
# 	columns =0
	
# 	p = document.add_paragraph()	
# 	p.add_run(str(number_q)+str('.')+worksheet.cell_value(quan_of_rows, columns)).bold = True
# 	number_q += 1

# 	columns+=1
# 	p.add_run()
# 	p = document.add_paragraph(str('{'))
# 	p.add_run(str('=')+str(worksheet.cell_value(quan_of_rows, columns))).font.highlight_color = WD_COLOR_INDEX.YELLOW	
	
# 	columns+=1
# 	p = document.add_paragraph(str('~')+str(worksheet.cell_value(quan_of_rows, columns)))
# 	columns+=1
# 	p = document.add_paragraph(str('~')+str(worksheet.cell_value(quan_of_rows, columns)))
# 	columns+=1
# 	p = document.add_paragraph(str('~')+str(worksheet.cell_value(quan_of_rows, columns)))
# 	p = document.add_paragraph(str('}'))
# 	columns+=1

# document.save('demoas.docx')



from aiogram import Bot, types
from aiogram.dispatcher import Dispatcher
from aiogram.utils import executor

from config import TOKEN

from aiogram.types import ReplyKeyboardRemove, \
    ReplyKeyboardMarkup, KeyboardButton, \
    InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.types import Message, CallbackQuery
from aiogram.types import CallbackQuery
import random


bot = Bot(token='6258839970:AAEIuBAdRLe3J7LZ57l3bcXm8y9yPmDHybY')
dp = Dispatcher(bot)
@dp.message_handler(commands=['start'])
async def process_start_command(message: types.Message):
    await message.reply("Are you ready?\n/yes")
# @dp.message_handler(commands=['yes'])
# async def process_start_command(message: types.Message):
#    pass 


@dp.message_handler(commands=['yes'])
async def url_command(message: types.Message):
    print('prea')
    variants = [1,2,3,4]
    random_var = random.shuffle(variants)
    random_ques = random.randint(0,500)
    
    first_var = worksheet.cell_value(random_ques, 1) 
    second_var = worksheet.cell_value(random_ques, 2)
    third_var = worksheet.cell_value(random_ques, 3)
    fourth_var = worksheet.cell_value(random_ques, 4)
    new_list = [(first_var, 'button_1_pressed' ),
                (second_var, 'button_2_pressed'),
                (third_var,'button_3_pressed'),
                (fourth_var, 'button_4_pressed')]
    random.shuffle(new_list)
    # urlkb = InlineKeyboardMarkup(row_width=1)
    urlkb = InlineKeyboardMarkup()
    # Button = InlineKeyboardButton(text='Перейти в блог Skillbox', 
    #     callback_data='button_1_pressed')
    # Button2 = InlineKeyboardButton(text='Перейти к курсам Skillbox',
    #     callback_data='button_2_pressed')
    # Button3 = InlineKeyboardButton(text='Перейти к курсам Skillbox',
    #     callback_data='button_3_pressed')
    # Button4 = InlineKeyboardButton(text='Перейти к курсам Skillbox',
    #     callback_data='button_4_pressed')
    text_variants = '===\n'

    for var, kol in new_list:

        urlkb.add(InlineKeyboardButton(text=var,
        callback_data=kol ))
        text_variants +=str(var)
        text_variants += '\n===\n'
    await message.reply(str(f"<b>{worksheet.cell_value(random_ques, 0)}</b>\n\n{text_variants}"), 
        reply_markup=urlkb, parse_mode='HTML')
 

@dp.callback_query_handler(text=['button_1_pressed'])
# async def process_button_1_press(callback: CallbackQuery):
#     
#     await callback.message.edit_text(
#         text='Alright\n',
#         reply_markup=callback.message.reply_markup, parse_mode='HTML')
async def sdfsdf(message: types.Message):
    await bot.send_message(message.from_user.id,text='<b>Alright</b>', parse_mode='HTML')

    random_ques = random.randint(0,500)
    
    first_var = worksheet.cell_value(random_ques, 1) 
    second_var = worksheet.cell_value(random_ques, 2)
    third_var = worksheet.cell_value(random_ques, 3)
    fourth_var = worksheet.cell_value(random_ques, 4)
    new_list = [(first_var, 'button_1_pressed' ),
                (second_var, 'button_2_pressed'),
                (third_var,'button_3_pressed'),
                (fourth_var, 'button_4_pressed')]
    random.shuffle(new_list)

    urlkb = InlineKeyboardMarkup()

    text_variants = '===\n'

    for var, kol in new_list:

        urlkb.add(InlineKeyboardButton(text=var,
        callback_data=kol ))
        text_variants +=str(var)
        text_variants += '\n===\n'
    await bot.send_message(message.from_user.id, str(f"<b>{worksheet.cell_value(random_ques, 0)}</b>\n\n{text_variants}"), 
        reply_markup=urlkb, parse_mode='HTML')


@dp.callback_query_handler(text=['button_2_pressed'])
async def btn2(message: types.Message):
    await message.answer("No")
@dp.callback_query_handler(text=['button_3_pressed'])
async def btn3(message: types.Message):
    await message.answer("No0")
@dp.callback_query_handler(text=['button_4_pressed'])
async def btn4(message: types.Message):
    await message.answer("Noooo")
if __name__ == '__main__':
    executor.start_polling(dp)
